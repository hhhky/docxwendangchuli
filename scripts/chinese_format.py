"""
╔══════════════════════════════════════════════════════════════╗
║     中文文档格式化核心模块  (Chinese Formatting Core)         ║
║     基于 python-docx · 支持中英文混排 · 标准排版规范          ║
╚══════════════════════════════════════════════════════════════╝

提供：
  - 字体预设（黑体/宋体/楷体/仿宋 + Times New Roman/Arial）
  - 页面设置预设（默认 A4 / 公文标准 / 学术论文）
  - 段落格式预设（行距、首行缩进、对齐方式）
  - 底层辅助函数（set_run_font、apply_paragraph_format）
  - 高层样式函数（style_title、style_heading、style_body、style_caption）

用法：
  from scripts.chinese_format import FontPreset, PAGE_DEFAULT, ChineseDocument

字体层级规则（Heiti/Songti Rule）：
  - 标题/各级标题 → 黑体 (Heiti) + Times New Roman
  - 正文 → 宋体 (Songti) + Times New Roman
  - 题注/表格 → 楷体 (Kaiti) + Arial
  - 公文正文 → 仿宋 (FangSong) + Times New Roman (GB/T 9704)

参考标准：
  - GB/T 9704-2012《党政机关公文格式》
  - GB/T 7713.2-2022《学术论文编写规则》
"""

from dataclasses import dataclass
from enum import Enum
from typing import Optional

from docx.shared import Pt, Cm, Inches, Emu
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor


# ═══════════════════════════════════════════════════════════════
# 字体规范定义
# ═══════════════════════════════════════════════════════════════

@dataclass(frozen=True)
class ChineseFontSet:
    """中英文混排字体组合。

    Attributes:
        cn: 中文字体名称（eastAsia）
        en: 西文字体名称（ascii / hAnsi）
        size_pt: 字号（磅值）
        bold: 是否加粗
        italic: 是否斜体
    """
    cn: str
    en: str
    size_pt: float
    bold: bool = False
    italic: bool = False


class FontPreset:
    """预定义字体方案。

    字号对照：
      一号 26pt   小二 18pt   四号 14pt   小五 9pt
      小一 24pt   三号 16pt   小四 12pt
      二号 22pt   小三 15pt   五号 10.5pt
    """

    # ── 标题系列（黑体） ──
    TITLE_LARGE = ChineseFontSet("黑体", "Times New Roman", 26, bold=True)   # 一号
    TITLE       = ChineseFontSet("黑体", "Times New Roman", 22, bold=True)   # 二号
    H1          = ChineseFontSet("黑体", "Times New Roman", 16, bold=True)   # 三号
    H2          = ChineseFontSet("黑体", "Times New Roman", 15, bold=True)   # 小三
    H3          = ChineseFontSet("黑体", "Times New Roman", 12, bold=True)   # 小四

    # ── 正文系列（宋体） ──
    BODY        = ChineseFontSet("宋体", "Times New Roman", 12, bold=False)   # 小四
    BODY_SMALL  = ChineseFontSet("宋体", "Times New Roman", 10.5, bold=False) # 五号

    # ── 题注系列（楷体） ──
    CAPTION     = ChineseFontSet("楷体", "Arial", 10.5, bold=False)           # 五号
    CAPTION_SM  = ChineseFontSet("楷体", "Arial", 9, bold=False)              # 小五

    # ── 公文专用（仿宋 / GB/T 9704） ──
    OFFICIAL_TITLE = ChineseFontSet("黑体", "Times New Roman", 22, bold=True)  # 公文标题用黑体
    OFFICIAL_BODY  = ChineseFontSet("仿宋", "Times New Roman", 16, bold=False) # 公文正文三号仿宋
    OFFICIAL_H1    = ChineseFontSet("黑体", "Times New Roman", 16, bold=True)  # 一级标题
    OFFICIAL_H2    = ChineseFontSet("楷体", "Times New Roman", 16, bold=True)  # 二级标题
    OFFICIAL_PAGENUM = ChineseFontSet("宋体", "Times New Roman", 14, bold=False) # 页码

    # ── 签发人/发文字号等 ──
    OFFICIAL_META  = ChineseFontSet("仿宋", "Times New Roman", 16, bold=False)


# ═══════════════════════════════════════════════════════════════
# 页面设置定义
# ═══════════════════════════════════════════════════════════════

@dataclass(frozen=True)
class ChinesePageSetup:
    """页面设置参数（单位：厘米）。

    A4 标准尺寸：21.0cm × 29.7cm
    """
    top: float
    bottom: float
    left: float
    right: float

    @property
    def content_width(self) -> float:
        """版心宽度（A4 宽度 - 左右边距）"""
        return 21.0 - self.left - self.right

    @property
    def content_height(self) -> float:
        """版心高度（A4 高度 - 上下边距）"""
        return 29.7 - self.top - self.bottom


# 默认页面设置（Word 标准边距）
PAGE_DEFAULT = ChinesePageSetup(top=2.54, bottom=2.54, left=3.18, right=3.18)

# 党政机关公文页面设置（GB/T 9704-2012）
PAGE_OFFICIAL = ChinesePageSetup(top=3.7, bottom=3.5, left=2.8, right=2.6)

# 学术论文页面设置（GB/T 7713）
PAGE_ACADEMIC = ChinesePageSetup(top=2.5, bottom=2.5, left=2.5, right=2.5)

# 工作总结/报告页面设置
PAGE_REPORT = ChinesePageSetup(top=3.0, bottom=2.5, left=2.8, right=2.6)


# ═══════════════════════════════════════════════════════════════
# 段落格式预设
# ═══════════════════════════════════════════════════════════════

@dataclass(frozen=True)
class ParagraphPreset:
    """段落格式预设。

    Attributes:
        line_spacing: 行距倍数（1.0=单倍, 1.25, 1.5, 2.0）
        space_before: 段前间距（磅值）
        space_after: 段后间距（磅值）
        first_line_indent_chars: 首行缩进字符数（0=不缩进, 2=标准正文）
        alignment: 对齐方式（WD_ALIGN_PARAGRAPH 或 None=默认）
    """
    line_spacing: float = 1.25
    space_before: float = 0
    space_after: float = 0
    first_line_indent_chars: int = 2
    alignment: Optional[int] = None


# 正文段落（首行缩进2字符，1.25倍行距，两端对齐）
PARA_BODY = ParagraphPreset(
    line_spacing=1.25,
    first_line_indent_chars=2,
    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
)

# 标题段落（无缩进，居中）
PARA_TITLE = ParagraphPreset(
    line_spacing=1.5,
    first_line_indent_chars=0,
    alignment=WD_ALIGN_PARAGRAPH.CENTER,
)

# 各级标题段落（无缩进，左对齐）
PARA_HEADING = ParagraphPreset(
    line_spacing=1.25,
    first_line_indent_chars=0,
)

# 题注段落（无缩进，居中）
PARA_CAPTION = ParagraphPreset(
    line_spacing=1.0,
    first_line_indent_chars=0,
    alignment=WD_ALIGN_PARAGRAPH.CENTER,
)

# 公文正文（固定28磅行距，首行缩进2字符）
PARA_OFFICIAL = ParagraphPreset(
    line_spacing=28.0 / 12.0,  # 固定28磅 ≈ 2.33倍行距（相对于12pt字）
    first_line_indent_chars=2,
    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
)

# 落款段落（无缩进，右对齐）
PARA_SIGNATURE = ParagraphPreset(
    line_spacing=1.5,
    first_line_indent_chars=0,
    alignment=WD_ALIGN_PARAGRAPH.RIGHT,
)


# ═══════════════════════════════════════════════════════════════
# 中文序号映射
# ═══════════════════════════════════════════════════════════════

CN_NUMBER = {
    1: "一", 2: "二", 3: "三", 4: "四", 5: "五",
    6: "六", 7: "七", 8: "八", 9: "九", 10: "十",
    11: "十一", 12: "十二", 13: "十三", 14: "十四", 15: "十五",
    16: "十六", 17: "十七", 18: "十八", 19: "十九", 20: "二十",
}


def cn_num(n: int) -> str:
    """将整数转换为中文数字（支持 1-99）。"""
    if n <= 0:
        return str(n)
    if n in CN_NUMBER:
        return CN_NUMBER[n]
    if n < 100:
        tens = n // 10
        ones = n % 10
        prefix = CN_NUMBER.get(tens, str(tens))
        if ones == 0:
            return f"{prefix}十"
        return f"{prefix}十{CN_NUMBER[ones]}"
    return str(n)


# ═══════════════════════════════════════════════════════════════
# 底层辅助函数
# ═══════════════════════════════════════════════════════════════

def set_run_font(
    run,
    cn: str = "宋体",
    en: str = "Times New Roman",
    size_pt: float = 12,
    bold: bool = False,
    italic: bool = False,
    color: Optional[str] = None,
):
    """为 Run 设置中英文字体（通过底层 XML 操作确保 eastAsia 字体生效）。

    python-docx 默认只设置 Western 字体（w:ascii / w:hAnsi），
    东亚字体（w:eastAsia）必须通过 XML 直接设置才能正确渲染中文。

    Args:
        run: python-docx Run 对象
        cn: 中文字体名称（如 "宋体"、"黑体"、"楷体"、"仿宋"）
        en: 西文字体名称（如 "Times New Roman"、"Arial"）
        size_pt: 字号（磅值）
        bold: 加粗
        italic: 斜体
        color: 字体颜色（可选，如 "FF0000"）
    """
    run.font.name = en
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size_pt)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)

    # 在 XML 层面设置东亚字体
    rPr = run._r.get_or_add_rPr()

    # 移除已有的 rFonts 元素
    existing = rPr.findall(qn('w:rFonts'))
    for ex in existing:
        rPr.remove(ex)

    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), en)       # ASCII/Latin 字体
    rFonts.set(qn('w:hAnsi'), en)       # 高 ANSI 字体
    rFonts.set(qn('w:eastAsia'), cn)    # 东亚字体
    rFonts.set(qn('w:cs'), en)          # 复杂脚本字体
    rPr.insert(0, rFonts)


def apply_paragraph_format(
    paragraph,
    preset: ParagraphPreset = PARA_BODY,
):
    """应用标准中文段落格式。

    Args:
        paragraph: python-docx Paragraph 对象
        preset: 段落格式预设（ParagraphPreset 实例）
    """
    pf = paragraph.paragraph_format
    pf.line_spacing = preset.line_spacing
    pf.space_before = Pt(preset.space_before)
    pf.space_after = Pt(preset.space_after)

    if preset.first_line_indent_chars > 0:
        # 首行缩进 = 字号 × 字符数（中文字符为正方形）
        # 默认基于小四号(12pt)计算
        pf.first_line_indent = Pt(12 * preset.first_line_indent_chars)

    if preset.alignment is not None:
        pf.alignment = preset.alignment


# ═══════════════════════════════════════════════════════════════
# 高层样式函数
# ═══════════════════════════════════════════════════════════════

def style_title(
    paragraph,
    text: str,
    font: ChineseFontSet = FontPreset.TITLE,
):
    """应用文档标题样式（居中、加粗、大号黑体）。

    Args:
        paragraph: python-docx Paragraph 对象
        text: 标题文字
        font: 字体预设（默认二号黑体）
    """
    apply_paragraph_format(paragraph, PARA_TITLE)
    paragraph.clear()
    run = paragraph.add_run(text)
    set_run_font(run, font.cn, font.en, font.size_pt, font.bold, font.italic)


def style_heading(
    paragraph,
    text: str,
    level: int = 1,
):
    """应用标题样式（左对齐、黑体系列）。

    Args:
        paragraph: python-docx Paragraph 对象
        text: 标题文字
        level: 标题级别（1/2/3，对应三号/小三/小四黑体）
    """
    fonts = {1: FontPreset.H1, 2: FontPreset.H2, 3: FontPreset.H3}
    font = fonts.get(level, FontPreset.H3)

    apply_paragraph_format(paragraph, PARA_HEADING)
    paragraph.clear()
    run = paragraph.add_run(text)
    set_run_font(run, font.cn, font.en, font.size_pt, font.bold)


def style_body(
    paragraph,
    text: str,
    font: ChineseFontSet = FontPreset.BODY,
):
    """应用正文样式（首行缩进2字符，两端对齐，宋体）。

    Args:
        paragraph: python-docx Paragraph 对象
        text: 正文文字
        font: 字体预设（默认小四宋体）
    """
    indent = 2 if font.size_pt >= 10 else 2
    preset = ParagraphPreset(
        line_spacing=1.25,
        first_line_indent_chars=indent,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )
    apply_paragraph_format(paragraph, preset)
    paragraph.clear()
    run = paragraph.add_run(text)
    set_run_font(run, font.cn, font.en, font.size_pt, font.bold)


def style_caption(
    paragraph,
    text: str,
    font: ChineseFontSet = FontPreset.CAPTION,
):
    """应用题注样式（居中、楷体）。

    Args:
        paragraph: python-docx Paragraph 对象
        text: 题注文字
        font: 字体预设（默认五号楷体）
    """
    apply_paragraph_format(paragraph, PARA_CAPTION)
    paragraph.clear()
    run = paragraph.add_run(text)
    set_run_font(run, font.cn, font.en, font.size_pt, font.bold)


def style_signature(
    paragraph,
    text: str,
    font: ChineseFontSet = FontPreset.BODY,
):
    """应用落款样式（右对齐、无缩进）。

    Args:
        paragraph: python-docx Paragraph 对象
        text: 落款文字
        font: 字体预设
    """
    apply_paragraph_format(paragraph, PARA_SIGNATURE)
    paragraph.clear()
    run = paragraph.add_run(text)
    set_run_font(run, font.cn, font.en, font.size_pt, font.bold)


def add_blank_line(document, count: int = 1, font: ChineseFontSet = FontPreset.BODY):
    """插入空白段落。

    Args:
        document: python-docx Document 对象
        count: 空白行数
        font: 字体预设（用于保持行距一致）
    """
    for _ in range(count):
        p = document.add_paragraph()
        apply_paragraph_format(p, ParagraphPreset(
            line_spacing=1.25,
            first_line_indent_chars=0,
        ))


def add_mixed_run(
    paragraph,
    segments: list[dict],
    clear: bool = False,
):
    """向段落添加混合格式的文本段。

    每个 segment 格式：
        {"text": "...", "cn": "宋体", "en": "Times New Roman",
         "size_pt": 12, "bold": False, "italic": False}

    Args:
        paragraph: python-docx Paragraph 对象
        segments: 文本段列表
        clear: 是否先清空段落
    """
    if clear:
        paragraph.clear()
    for seg in segments:
        run = paragraph.add_run(seg.get("text", ""))
        set_run_font(
            run,
            cn=seg.get("cn", "宋体"),
            en=seg.get("en", "Times New Roman"),
            size_pt=seg.get("size_pt", 12),
            bold=seg.get("bold", False),
            italic=seg.get("italic", False),
        )


# ═══════════════════════════════════════════════════════════════
# 页面设置应用函数
# ═══════════════════════════════════════════════════════════════

def apply_page_setup(document, page_setup: ChinesePageSetup):
    """应用页面设置到文档的第一节。

    Args:
        document: python-docx Document 对象
        page_setup: 页面设置预设
    """
    section = document.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(page_setup.top)
    section.bottom_margin = Cm(page_setup.bottom)
    section.left_margin = Cm(page_setup.left)
    section.right_margin = Cm(page_setup.right)
