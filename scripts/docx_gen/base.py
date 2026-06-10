"""
ChineseDocument 工厂类

提供创建中文格式化 Word 文档的统一入口。
所有模板函数均基于此类构建。
"""

import os
from docx import Document
from docx.shared import Pt, Cm

# 当作为 skill 脚本调用时，使用相对导入
try:
    from ..chinese_format import (
        FontPreset, ChineseFontSet, ChinesePageSetup, ParagraphPreset,
        PAGE_DEFAULT, PAGE_OFFICIAL, PAGE_ACADEMIC, PAGE_REPORT,
        PARA_BODY, PARA_TITLE, PARA_HEADING, PARA_CAPTION, PARA_SIGNATURE, PARA_OFFICIAL,
        set_run_font, apply_paragraph_format, apply_page_setup,
        style_title, style_heading, style_body, style_caption, style_signature,
        add_blank_line,
    )
except ImportError:
    # 回退：直接导入（从 scripts/ 目录运行时）
    from chinese_format import (
        FontPreset, ChineseFontSet, ChinesePageSetup, ParagraphPreset,
        PAGE_DEFAULT, PAGE_OFFICIAL, PAGE_ACADEMIC, PAGE_REPORT,
        PARA_BODY, PARA_TITLE, PARA_HEADING, PARA_CAPTION, PARA_SIGNATURE, PARA_OFFICIAL,
        set_run_font, apply_paragraph_format, apply_page_setup,
        style_title, style_heading, style_body, style_caption, style_signature,
        add_blank_line,
    )


class ChineseDocument:
    """中文 Word 文档工厂类。

    封装了文档创建、页面设置、元素添加、保存等操作，
    确保所有文档具有一致的中文排版规范。

    用法：
        doc = ChineseDocument.create(PAGE_DEFAULT)
        ChineseDocument.add_title(doc, "我的文档")
        ChineseDocument.add_body(doc, "正文内容...")
        ChineseDocument.save(doc, "output.docx")
    """

    # ── 文档创建 ──────────────────────────────────────────

    @staticmethod
    def create(page_setup: ChinesePageSetup = PAGE_DEFAULT) -> Document:
        """创建空白 A4 文档并应用页面设置。

        Args:
            page_setup: 页面设置预设（默认 PAGE_DEFAULT）

        Returns:
            python-docx Document 对象
        """
        doc = Document()
        apply_page_setup(doc, page_setup)
        return doc

    # ── 元素添加 ──────────────────────────────────────────

    @staticmethod
    def add_title(doc: Document, text: str, large: bool = False) -> "Paragraph":
        """添加文档主标题（居中、黑体、加粗）。

        Args:
            doc: Document 对象
            text: 标题文字
            large: 是否使用一号字（26pt），默认二号（22pt）

        Returns:
            创建的 Paragraph 对象
        """
        font = FontPreset.TITLE_LARGE if large else FontPreset.TITLE
        p = doc.add_paragraph()
        style_title(p, text, font)
        return p

    @staticmethod
    def add_heading(doc: Document, text: str, level: int = 1) -> "Paragraph":
        """添加章节标题（左对齐、黑体系列）。

        Args:
            doc: Document 对象
            text: 标题文字
            level: 标题级别 1/2/3

        Returns:
            创建的 Paragraph 对象
        """
        p = doc.add_paragraph()
        style_heading(p, text, level)
        return p

    @staticmethod
    def add_body(doc: Document, text: str, font: ChineseFontSet = None) -> "Paragraph":
        """添加正文段落（首行缩进2字符、两端对齐、宋体）。

        Args:
            doc: Document 对象
            text: 正文文字
            font: 字体预设（默认小四宋体）

        Returns:
            创建的 Paragraph 对象
        """
        f = font or FontPreset.BODY
        p = doc.add_paragraph()
        style_body(p, text, f)
        return p

    @staticmethod
    def add_caption(doc: Document, text: str) -> "Paragraph":
        """添加题注（居中、楷体）。

        Args:
            doc: Document 对象
            text: 题注文字

        Returns:
            创建的 Paragraph 对象
        """
        p = doc.add_paragraph()
        style_caption(p, text)
        return p

    @staticmethod
    def add_signature(doc: Document, text: str) -> "Paragraph":
        """添加落款（右对齐、无缩进）。

        Args:
            doc: Document 对象
            text: 落款文字

        Returns:
            创建的 Paragraph 对象
        """
        p = doc.add_paragraph()
        style_signature(p, text)
        return p

    @staticmethod
    def add_blank(doc: Document, count: int = 1):
        """插入空白行。

        Args:
            doc: Document 对象
            count: 行数
        """
        add_blank_line(doc, count)

    # ── 保存 ──────────────────────────────────────────────

    @staticmethod
    def save(doc: Document, path: str) -> str:
        """保存文档到指定路径。

        Args:
            doc: Document 对象
            path: 输出路径（.docx）

        Returns:
            输出路径
        """
        # 确保目录存在
        os.makedirs(os.path.dirname(os.path.abspath(path)), exist_ok=True)
        doc.save(path)
        return path
